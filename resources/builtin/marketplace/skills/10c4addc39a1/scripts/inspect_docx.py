#!/usr/bin/env python3
from __future__ import annotations
import argparse, json, re, zipfile
from pathlib import Path
from docx import Document

REQUIRED_HEADINGS=[
'执行摘要','目录','1. 任务定位、政策语境与研究边界','2. 方法框架：共享上层本体、四层模型与证据等级',
'2.4 网站检索与来源核验门禁','3. ','4. 关键参与方与角色','5. 主要业务系统与 System of Record','6. Canonical Concepts 基线',
'7. 核心关系三元组','8. 端到端核心流程','9. 事件、决策、规则、风险与 KPI',
'10. 外部标准、政策语义资产与本体对标','11. Shared Upper Ontology 映射',
'12. ','13. 闭环示例与 KSTAR/ECS 映射','14. 架构边界、实施建议与待确认问题',
'附录 A. Master Mapping Workbook 字段与示例','附录 B. 来源、版本与证据清单','B.1 网站检索与来源核验摘要','B.2 报告来源台账','附录 C. 来源材料与分析边界说明']

def main():
    ap=argparse.ArgumentParser(); ap.add_argument('docx'); ap.add_argument('--data'); ap.add_argument('--out'); args=ap.parse_args()
    p=Path(args.docx); d=Document(p)
    headings=[x.text.strip() for x in d.paragraphs if x.style and x.style.name.startswith('Heading')]
    errors=[]; warnings=[]
    for h in REQUIRED_HEADINGS:
        if h.endswith('. '):
            if not any(x.startswith(h) for x in headings): errors.append(f'missing heading prefix: {h}')
        elif h not in headings: errors.append(f'missing heading: {h}')
    if len(d.tables)<20: warnings.append(f'low table count: {len(d.tables)}')
    # OOXML checks.
    with zipfile.ZipFile(p) as z:
        names=set(z.namelist()); xml=z.read('word/document.xml').decode('utf-8','ignore')
        image_count=len([n for n in names if n.startswith('word/media/')])
        toc='TOC \\o' in xml
        page_field=' PAGE ' in ''.join(z.read(n).decode('utf-8','ignore') for n in names if n.startswith('word/header') or n.startswith('word/footer'))
        alt_count=len(re.findall(r' descr="[^"]+"',xml))
        if image_count<3: errors.append(f'expected at least 3 images, found {image_count}')
        if not toc: errors.append('Word TOC field missing')
        if not page_field: warnings.append('page-number field not detected')
        if alt_count<image_count: warnings.append(f'alt text count {alt_count} < images {image_count}')
    counts={"paragraphs":len(d.paragraphs),"headings":len(headings),"tables":len(d.tables),"images":image_count,"alt_text":alt_count}
    if args.data:
        data=json.load(open(args.data,encoding='utf-8'))
        expected={k:len(data.get(k,[])) for k in ['concepts','relations','processes','standards','sources']}
        counts.update(expected)
        # crude text-presence checks.
        full='\n'.join(x.text for x in d.paragraphs)
        for c in data.get('concepts',[])[:5]:
            if c['name_cn'] not in full and not any(c['name_cn'] in cell.text for t in d.tables for row in t.rows for cell in row.cells): warnings.append(f"concept not detected in document: {c['name_cn']}")
        all_cells='\n'.join(cell.text for t in d.tables for row in t.rows for cell in row.cells)
        if data.get('research_assurance',{}).get('gate_status')!='passed': errors.append('research assurance not passed in report data')
        if '研究门状态' not in all_cells or 'passed' not in all_cells: errors.append('research-gate summary not detected in DOCX')
    score=max(0,100-len(errors)*12-len(warnings)*2)
    report={"status":"passed" if not errors else "failed","docx":str(p),"counts":counts,"errors":errors,"warnings":warnings,"score":score}
    out=Path(args.out) if args.out else p.with_name(p.stem+'-verification.json'); out.write_text(json.dumps(report,ensure_ascii=False,indent=2),encoding='utf-8')
    print(json.dumps(report,ensure_ascii=False,indent=2)); raise SystemExit(0 if not errors else 1)
if __name__=='__main__': main()
