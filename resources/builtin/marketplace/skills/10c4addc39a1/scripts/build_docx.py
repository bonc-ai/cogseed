#!/usr/bin/env python3
from __future__ import annotations
import argparse, json, math, os, tempfile
from collections import defaultdict
from pathlib import Path
from typing import Any, Iterable, Sequence

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from matplotlib.font_manager import FontProperties

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

NAVY='17365D'; BLUE='2F75B5'; TEAL='1F7A8C'; GREEN='4F8A5B'; ORANGE='C55A11'; RED='A61B1B'; GOLD='BF9000'
LIGHT_BLUE='D9EAF7'; LIGHT_TEAL='DDEBF7'; LIGHT_GREEN='E2F0D9'; LIGHT_ORANGE='FCE4D6'; LIGHT_GRAY='F2F2F2'; WHITE='FFFFFF'; DARK='404040'
FONT_NAME='Noto Sans CJK SC'
FONT_REG='/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc'
FONT_BOLD='/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc'
font_reg=FontProperties(fname=FONT_REG) if Path(FONT_REG).exists() else None
font_bold=FontProperties(fname=FONT_BOLD) if Path(FONT_BOLD).exists() else font_reg

UPPER_MEANINGS={
'Party':'人、组织、部门、监管方、服务提供者或受益人','Role':'主体在具体情境中扮演的角色','Agent':'人、AI智能体、机器人、软件服务或组织单元','Goal':'期望达到的状态或业务目标','Situation':'问题、需要、机会、上下文或触发条件','Capability':'主体、系统或智能体能够且获授权执行的能力','Service/Product':'向受益方交付的价值','Process/Case':'为达成目标而组织的生命周期','Task/Action':'流程中的可执行工作单元','Resource/Asset':'资金、设备、设施、材料、数据、知识等资源','Information Object':'文档、模型、记录、消息、数据集或要求','Event':'发生并改变状态的事项','Decision':'基于证据和规则选择行动','Policy/Rule/Constraint':'必须、可以或不得发生的约束','Agreement/Commitment':'合同、订单、同意、许可、SLA等承诺','Observation/Evidence':'支持结论的观测或主张','Risk/Control':'不确定性、义务、危害及其缓解控制','Outcome':'行动或流程产生的结果','Measure/KPI':'对状态或结果的定量/定性评估','Episode/Learning':'情境—行动—结果—经验的可复用记录','System/Interface':'应用、服务、数据库、连接器、API、CLI或MCP通道'}

def rgb(h): return tuple(int(h[i:i+2],16)/255 for i in (0,2,4))

def load(path):
    with open(path,encoding='utf-8') as f: return json.load(f)

def joinv(v, sep='；'):
    if v is None: return ''
    if isinstance(v,list): return sep.join(str(x) for x in v)
    if isinstance(v,dict): return sep.join(f'{k}={joinv(x)}' for k,x in v.items())
    return str(v)

def source_text(v):
    if isinstance(v,dict):
        text=str(v.get('text','')); refs=v.get('source_refs',[])
        return text + (f"〔{'/'.join(refs)}〕" if refs else '')
    return str(v)

def set_run_font(run,size=10.5,bold=None,color=None,italic=None,name=FONT_NAME):
    run.font.name=name
    if run._element.get_or_add_rPr() is not None:
        run._element.rPr.rFonts.set(qn('w:eastAsia'),name)
        run._element.rPr.rFonts.set(qn('w:ascii'),name)
        run._element.rPr.rFonts.set(qn('w:hAnsi'),name)
    run.font.size=Pt(size)
    if bold is not None: run.bold=bold
    if color: run.font.color.rgb=RGBColor.from_string(color)
    if italic is not None: run.italic=italic

def set_cell_shading(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd'))
    if shd is None: shd=OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'),fill)

def set_cell_margins(cell, top=70,start=70,bottom=70,end=70):
    tcPr=cell._tc.get_or_add_tcPr(); tcMar=tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar=OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node=tcMar.find(qn(f'w:{m}'))
        if node is None: node=OxmlElement(f'w:{m}'); tcMar.append(node)
        node.set(qn('w:w'),str(v)); node.set(qn('w:type'),'dxa')

def repeat_header(row):
    trPr=row._tr.get_or_add_trPr(); e=OxmlElement('w:tblHeader'); e.set(qn('w:val'),'true'); trPr.append(e)

def cant_split(row):
    trPr=row._tr.get_or_add_trPr(); e=OxmlElement('w:cantSplit'); trPr.append(e)

def set_col_width(cell,cm):
    cell.width=Cm(cm); tcPr=cell._tc.get_or_add_tcPr(); tcW=tcPr.find(qn('w:tcW'))
    if tcW is None: tcW=OxmlElement('w:tcW'); tcPr.append(tcW)
    tcW.set(qn('w:w'),str(int(cm*567))); tcW.set(qn('w:type'),'dxa')

def add_page_number(p):
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(); b=OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'),'begin')
    instr=OxmlElement('w:instrText'); instr.set(qn('xml:space'),'preserve'); instr.text=' PAGE '
    e=OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'),'end')
    r._r.extend([b,instr,e]); set_run_font(r,size=9,color='666666')

def add_toc(doc):
    entries=[
      '执行摘要',
      '1. 任务定位、政策语境与研究边界',
      '2. 方法框架：共享上层本体、四层模型与证据等级',
      '3. 领域边界与模块划分',
      '4. 关键参与方与角色',
      '5. 主要业务系统与 System of Record',
      '6. Canonical Concepts 基线',
      '7. 核心关系三元组',
      '8. 端到端核心流程',
      '9. 事件、决策、规则、风险与 KPI',
      '10. 外部标准、政策语义资产与本体对标',
      '11. Shared Upper Ontology 映射',
      '12. 智能体扩展、授权、HITL与审计',
      '13. 闭环示例与 KSTAR/ECS 映射',
      '14. 架构边界、实施建议与待确认问题',
      '附录 A. Master Mapping Workbook',
      '附录 B. 来源、版本与证据',
      '附录 C. 来源材料与分析边界'
    ]
    for entry in entries:
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2.3); p.paragraph_format.left_indent=Cm(.35)
        r=p.add_run(entry); set_run_font(r,size=10.0,color=DARK)
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(5)
    r=p.add_run(); b=OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'),'begin')
    instr=OxmlElement('w:instrText'); instr.set(qn('xml:space'),'preserve'); instr.text=' TOC \\o "1-3" \\h \\z \\u '
    sep=OxmlElement('w:fldChar'); sep.set(qn('w:fldCharType'),'separate')
    t=OxmlElement('w:t'); t.text='（Word目录字段：在桌面版Word中更新后显示页码）'
    er=OxmlElement('w:r'); er.append(t); end=OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'),'end')
    r._r.extend([b,instr,sep]); p._p.append(er); r2=p.add_run(); r2._r.append(end)

def set_update_fields(doc):
    settings=doc.settings._element; u=settings.find(qn('w:updateFields'))
    if u is None: u=OxmlElement('w:updateFields'); settings.append(u)
    u.set(qn('w:val'),'true')

def add_heading(doc,text,level=1):
    p=doc.add_paragraph(style=f'Heading {level}'); p.paragraph_format.keep_with_next=True
    r=p.add_run(text); return p

def add_para(doc,text='',refs=None,bold_prefix=None,color=DARK,align=None,after=5,line=1.22):
    p=doc.add_paragraph();
    if align is not None: p.alignment=align
    if bold_prefix and text.startswith(bold_prefix):
        r=p.add_run(bold_prefix); set_run_font(r,bold=True,color=color)
        r=p.add_run(text[len(bold_prefix):]); set_run_font(r,color=color)
    else:
        r=p.add_run(text); set_run_font(r,color=color)
    if refs:
        rr=p.add_run(f"〔{'/'.join(refs)}〕"); set_run_font(rr,size=8.5,color=BLUE)
    p.paragraph_format.space_after=Pt(after); p.paragraph_format.line_spacing=line
    return p

def add_cited(doc,obj):
    if isinstance(obj,dict): return add_para(doc,str(obj.get('text','')),refs=obj.get('source_refs',[]))
    return add_para(doc,str(obj))

def add_bullet(doc,text,level=0,refs=None):
    p=doc.add_paragraph(style='List Bullet' if level==0 else 'List Bullet 2'); r=p.add_run(text); set_run_font(r,size=10.1)
    if refs: rr=p.add_run(f"〔{'/'.join(refs)}〕"); set_run_font(rr,size=8.3,color=BLUE)
    p.paragraph_format.space_after=Pt(2.5); p.paragraph_format.line_spacing=1.15; return p

def add_number(doc,text):
    p=doc.add_paragraph(style='List Number'); r=p.add_run(text); set_run_font(r,size=10.1); p.paragraph_format.space_after=Pt(2.5); return p

def add_callout(doc,title,text,fill=LIGHT_BLUE,edge=NAVY):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; repeat_header(t.rows[0]); c=t.cell(0,0); set_cell_shading(c,fill); set_cell_margins(c,110,130,110,130)
    p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); r=p.add_run(title+'  '); set_run_font(r,bold=True,color=edge,size=10.5); r=p.add_run(text); set_run_font(r,size=10.1)
    doc.add_paragraph().paragraph_format.space_after=Pt(1); return t

def add_table(doc,headers,rows,widths=None,font_size=8.3,header_fill=NAVY,first_col_fill=None,cell_margin=65,header_margin=75):
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    repeat_header(t.rows[0])
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; set_cell_shading(c,header_fill); set_cell_margins(c,header_margin,cell_margin,header_margin,cell_margin); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(0); r=p.add_run(str(h)); set_run_font(r,size=font_size,bold=True,color=WHITE)
        if widths: set_col_width(c,widths[i])
    for ri,row in enumerate(rows):
        cells=t.add_row().cells; cant_split(t.rows[-1])
        for i,v in enumerate(row):
            c=cells[i]; set_cell_margins(c,cell_margin,cell_margin,cell_margin,cell_margin); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP
            if widths: set_col_width(c,widths[i])
            if first_col_fill and i==0: set_cell_shading(c,first_col_fill)
            elif ri%2==1: set_cell_shading(c,'FAFAFA')
            p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.05
            r=p.add_run(joinv(v)); set_run_font(r,size=font_size,bold=bool(first_col_fill and i==0))
    doc.add_paragraph().paragraph_format.space_after=Pt(1); return t

def add_figure(doc,path,caption,width=6.45):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; shape=p.add_run().add_picture(str(path),width=Inches(width))
    shape._inline.docPr.set('descr',caption); shape._inline.docPr.set('title',caption)
    cp=doc.add_paragraph(); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=cp.add_run(caption); set_run_font(r,size=9,color='666666',italic=True); cp.paragraph_format.space_after=Pt(7)

def new_page(doc): doc.add_page_break()

def create_method_diagram(data,path):
    fig,ax=plt.subplots(figsize=(12,5.4),dpi=180); ax.set_xlim(0,12); ax.set_ylim(0,5.4); ax.axis('off')
    layers=data['methodology']['four_layers']; fills=[LIGHT_BLUE,LIGHT_TEAL,LIGHT_GRAY,LIGHT_GREEN]; edges=[NAVY,BLUE,DARK,GREEN]
    for i,l in enumerate(layers):
        y=4.3-i*1.02; box=FancyBboxPatch((0.45,y),5.2,.75,boxstyle='round,pad=.03',facecolor=rgb(fills[i]),edgecolor=rgb(edges[i]),linewidth=1.7); ax.add_patch(box)
        ax.text(.75,y+.48,l['name'],fontproperties=font_bold,fontsize=13,color=rgb(edges[i]),va='center'); ax.text(2.25,y+.35,joinv(l.get('examples',[])[:4],' / '),fontproperties=font_reg,fontsize=9.5,color=rgb(DARK),va='center')
    flow=data['methodology'].get('common_flow',[]); xs=[6.15+i*.73 for i in range(min(8,len(flow)))]
    for i,(x,label) in enumerate(zip(xs,flow[:8])):
        box=FancyBboxPatch((x,2.1),.62,1.2,boxstyle='round,pad=.02',facecolor=rgb(LIGHT_BLUE),edgecolor=rgb(BLUE),linewidth=1.2); ax.add_patch(box)
        ax.text(x+.31,2.7,str(label).replace('/','\n/'),fontproperties=font_bold,fontsize=6.7,ha='center',va='center',color=rgb(NAVY))
        if i<len(xs)-1: ax.add_patch(FancyArrowPatch((x+.63,2.7),(xs[i+1]-.01,2.7),arrowstyle='-|>',mutation_scale=9,color=rgb(BLUE),linewidth=1))
    ax.text(6.2,4.35,'共同业务闭环',fontproperties=font_bold,fontsize=16,color=rgb(NAVY))
    ax.text(6.2,.8,'纪律：现实对象 ≠ 系统记录；信号/推荐 ≠ 正式决定；\n证据、授权、HITL与审计贯穿智能体执行。',fontproperties=font_reg,fontsize=11,color=rgb(DARK))
    fig.tight_layout(); fig.savefig(path,bbox_inches='tight'); plt.close(fig)

def create_module_diagram(data,path):
    mods=data['modules']; n=len(mods); cols=3 if n>4 else 2; rows=math.ceil(n/cols)
    fig_h=2.1+rows*1.35; fig,ax=plt.subplots(figsize=(12,fig_h),dpi=180); ax.set_xlim(0,12); ax.set_ylim(0,fig_h); ax.axis('off')
    top_y=fig_h-.75; b=FancyBboxPatch((.7,top_y),10.6,.55,boxstyle='round,pad=.03',facecolor=rgb(LIGHT_BLUE),edgecolor=rgb(NAVY),linewidth=1.7); ax.add_patch(b)
    ax.text(6,top_y+.28,'Shared Upper Ontology',fontproperties=font_bold,fontsize=13,ha='center',va='center',color=rgb(NAVY))
    colors=[(LIGHT_TEAL,BLUE),(LIGHT_GREEN,GREEN),(LIGHT_ORANGE,ORANGE),(LIGHT_BLUE,NAVY),('FFF2CC',GOLD),(LIGHT_GRAY,DARK)]
    w=3.35 if cols==3 else 5.0; xgap=(12-cols*w)/(cols+1); y0=top_y-1.25
    for i,m in enumerate(mods):
        r=i//cols; c=i%cols; x=xgap+c*(w+xgap); y=y0-r*1.25; fill,edge=colors[i%len(colors)]
        box=FancyBboxPatch((x,y),w,.9,boxstyle='round,pad=.04',facecolor=rgb(fill),edgecolor=rgb(edge),linewidth=1.4); ax.add_patch(box)
        ax.text(x+w/2,y+.62,m['name'],fontproperties=font_bold,fontsize=10.5,ha='center',va='center',color=rgb(edge)); ax.text(x+w/2,y+.27,str(m['core_objects'])[:62],fontproperties=font_reg,fontsize=7.5,ha='center',va='center',color=rgb(DARK),wrap=True)
        ax.add_patch(FancyArrowPatch((6,top_y),(x+w/2,y+.92),arrowstyle='-|>',mutation_scale=8,color=rgb('8A8A8A'),linewidth=.8))
    ax.text(6,.25,'领域模块保持专业特化；系统字段与组织扩展通过映射连接，不反向污染核心语义。',fontproperties=font_bold,fontsize=10.5,ha='center',color=rgb(NAVY))
    fig.tight_layout(); fig.savefig(path,bbox_inches='tight'); plt.close(fig)

def create_loop_diagram(data,path):
    labels=data['methodology'].get('common_flow',[])[:8]; fills=[LIGHT_BLUE,LIGHT_TEAL,'EAF2F8','FFF2CC',LIGHT_GREEN,LIGHT_ORANGE,'E4DFEC',LIGHT_GRAY]; edges=[NAVY,BLUE,TEAL,GOLD,GREEN,ORANGE,'7030A0',DARK]
    fig,ax=plt.subplots(figsize=(12,4.8),dpi=180); ax.set_xlim(0,12); ax.set_ylim(0,4.8); ax.axis('off')
    coords=[(.4,2.65),(3.2,2.65),(6,2.65),(8.8,2.65),(8.8,.9),(6,.9),(3.2,.9),(.4,.9)]
    for i,(label,(x,y)) in enumerate(zip(labels,coords)):
        box=FancyBboxPatch((x,y),2.35,.85,boxstyle='round,pad=.03',facecolor=rgb(fills[i]),edgecolor=rgb(edges[i]),linewidth=1.4); ax.add_patch(box)
        ax.text(x+1.175,y+.43,str(label),fontproperties=font_bold,fontsize=9.5,ha='center',va='center',color=rgb(edges[i]))
        if i<len(labels)-1:
            nx,ny=coords[i+1]; ax.add_patch(FancyArrowPatch((x+1.175,y),(nx+1.175,ny+.86),arrowstyle='-|>',mutation_scale=10,color=rgb(BLUE),linewidth=1.0,connectionstyle='arc3,rad=0.0'))
    if labels: ax.add_patch(FancyArrowPatch((1.55,.9),(1.55,2.65),arrowstyle='-|>',mutation_scale=10,color=rgb(GREEN),linewidth=1.2,connectionstyle='arc3,rad=-.35'))
    ax.text(6,4.3,'Situation → Goal → Process → Task/Decision → Evidence/Event → Outcome/KPI → Episode/Learning',fontproperties=font_bold,fontsize=12,ha='center',color=rgb(NAVY))
    ax.text(6,.25,'KSTAR：预测 → 实际 → ΔA/ΔR → 归因 → 有界候选更新；自动状态最高 staged。',fontproperties=font_reg,fontsize=10.5,ha='center',color=rgb(DARK))
    fig.tight_layout(); fig.savefig(path,bbox_inches='tight'); plt.close(fig)

def configure_doc(doc,data):
    sec=doc.sections[0]; sec.page_width=Cm(21); sec.page_height=Cm(29.7); sec.top_margin=Cm(1.75); sec.bottom_margin=Cm(1.65); sec.left_margin=Cm(1.8); sec.right_margin=Cm(1.8)
    styles=doc.styles
    normal=styles['Normal']; normal.font.name=FONT_NAME; normal._element.rPr.rFonts.set(qn('w:eastAsia'),FONT_NAME); normal.font.size=Pt(10.5)
    for level,size,color in [(1,16,NAVY),(2,13.2,BLUE),(3,11.2,TEAL)]:
        s=styles[f'Heading {level}']; s.font.name=FONT_NAME; s._element.rPr.rFonts.set(qn('w:eastAsia'),FONT_NAME); s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color); s.paragraph_format.space_before=Pt(10 if level==1 else 7); s.paragraph_format.space_after=Pt(5); s.paragraph_format.keep_with_next=True
    # Header/footer all sections.
    for s in doc.sections:
        hp=s.header.paragraphs[0]; hp.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=hp.add_run(f"第{data['meta']['volume_number']}项 · {data['meta']['domain_name_cn']} · Ontology分析与对标 · {data['meta']['report_version']}"); set_run_font(r,size=8.5,color='666666')
        fp=s.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=fp.add_run('候选稿 / staged for review  ·  '); set_run_font(r,size=8.5,color='777777'); add_page_number(fp)
    set_update_fields(doc)
    cp=doc.core_properties; cp.title=data['meta']['title']; cp.subject='Domain Ontology Landscape and Alignment Report'; cp.author='Ontology Landscape & Alignment Report Skill'; cp.keywords=', '.join(data['meta'].get('keywords',[])); cp.comments='Staged candidate; not a complete production ontology or official standard.'

def build(data,out_path):
    out=Path(out_path); out.parent.mkdir(parents=True,exist_ok=True); asset=out.parent/(out.stem+'_assets'); asset.mkdir(parents=True,exist_ok=True)
    method=asset/'method.png'; modules=asset/'modules.png'; loop=asset/'loop.png'; create_method_diagram(data,method); create_module_diagram(data,modules); create_loop_diagram(data,loop)
    doc=Document(); configure_doc(doc,data)
    # Cover
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(75); r=p.add_run(data['meta']['title']); set_run_font(r,size=25,bold=True,color=NAVY)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(data['meta'].get('subtitle','Domain Ontology Landscape and Alignment Report')); set_run_font(r,size=13,color=BLUE)
    doc.add_paragraph();
    add_callout(doc,'交付定位',data['meta'].get('scope_statement','领域本体景观与对齐地图，不是完整领域本体实现。'),fill=LIGHT_BLUE,edge=NAVY)
    for label,val in [('分册',f"第{data['meta']['volume_number']}项 / {data['meta']['domain_name_cn']}"),('版本',data['meta']['report_version']),('研究时点',data['meta']['research_date']),('交付对象',data['meta'].get('deliverable_for','')),('状态',data['meta'].get('status','candidate')),('保密标识',data['meta'].get('confidentiality','内部研究候选稿'))]:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(f'{label}：'); set_run_font(r,size=10,bold=True,color=NAVY); r=p.add_run(str(val)); set_run_font(r,size=10,color=DARK)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(35); r=p.add_run('自动化流程最高状态：staged；最终内容、专业边界与发布由人类负责人决定。'); set_run_font(r,size=9.5,color=RED,italic=True)
    new_page(doc)
    # Executive summary
    add_heading(doc,'执行摘要',1)
    for para in data['executive_summary']['paragraphs']: add_cited(doc,para)
    add_heading(doc,'核心结论',2)
    for c in data['executive_summary']['core_conclusions']:
        add_callout(doc,c['title'],c['text']+(f"〔{'/'.join(c.get('source_refs',[]))}〕" if c.get('source_refs') else ''),fill=LIGHT_GREEN,edge=GREEN)
    ra=data.get('research_assurance',{})
    counts=[('Canonical Concepts',len(data['concepts'])),('核心关系',len(data['relations'])),('端到端流程',len(data['processes'])),('标准/语义资产',len(data['standards'])),('参与方/角色',len(data['stakeholders'])),('系统/SoR',len(data['systems'])),('闭环示例',len(data['closed_loops'])),('来源记录',len(data['sources'])),('网站检索查询',ra.get('queries_executed',0)),('已打开来源',ra.get('opened_sources',0)),('已核验一手来源',ra.get('verified_primary_sources',0))]
    add_heading(doc,'成果规模',2); add_table(doc,['内容','数量'],counts,[10.5,3.0],font_size=9.3,first_col_fill=LIGHT_BLUE)
    add_callout(doc,'证据边界',data['meta'].get('research_note','报告区分来源事实、外部核实、分析建议和待确认事项。'),fill=LIGHT_ORANGE,edge=ORANGE)
    new_page(doc)
    add_heading(doc,'目录',1); add_toc(doc); add_para(doc,'提示：在桌面版 Word 中打开后更新目录和页码字段。',color='666666')

    # 1
    new_page(doc); add_heading(doc,'1. 任务定位、政策语境与研究边界',1)
    add_heading(doc,'1.1 任务目标',2); add_cited(doc,data['policy_context']['task_objective'])
    add_heading(doc,'1.2 政策场景定义与治理约束',2)
    for x in data['policy_context']['policy_scope']: add_cited(doc,x)
    add_heading(doc,'1.3 纳入范围与排除范围',2)
    add_table(doc,['纳入范围','排除范围'],[[joinv(data['policy_context']['included_scope'],'\n'),joinv(data['policy_context']['excluded_scope'],'\n')]],[8.4,8.4],font_size=9)
    add_heading(doc,'1.4 边界澄清与建模纪律',2)
    for x in data['policy_context']['boundary_clarifications']: add_bullet(doc,x)
    for x in data.get('semantic_confusions',[]): add_callout(doc,x['expression'],x['clarification'],fill=LIGHT_ORANGE,edge=ORANGE)

    # 2
    new_page(doc); add_heading(doc,'2. 方法框架：共享上层本体、四层模型与证据等级',1)
    add_heading(doc,'2.1 Shared Upper Ontology',2)
    umap={x['upper_construct']:x for x in data['upper_mappings']}
    rows=[]
    for u in data['methodology']['shared_upper_ontology']:
        rows.append([u,UPPER_MEANINGS.get(u,''),joinv(umap.get(u,{}).get('domain_concepts',[]),'、')])
    add_table(doc,['Shared Construct','通用含义','本领域示例'],rows,[4.2,7.3,5.3],font_size=8.2)
    add_heading(doc,'2.2 四层建模纪律',2)
    add_table(doc,['层次','定义','领域示例'],[[x['name'],x['description'],joinv(x.get('examples',[]),'、')] for x in data['methodology']['four_layers']],[3.2,7.0,6.6],font_size=8.7,first_col_fill=LIGHT_BLUE)
    add_figure(doc,method,'图 2-1  统一上层本体、四层语义与共同业务闭环')
    add_heading(doc,'2.3 证据和成熟度',2)
    add_table(doc,['代码','含义','使用方式'],[[x['code'],x['meaning'],x['usage']] for x in data['methodology']['evidence_levels']],[2.2,5.2,9.4],font_size=8.8,first_col_fill=LIGHT_GREEN)
    add_callout(doc,'成熟度', 'Canonical Concept、关系、规则和映射应标注 candidate / reviewed / approved / deprecated；未经专家确认的R-Box理由和阈值不得写成正式规则。',fill=LIGHT_GREEN,edge=GREEN)
    new_page(doc); add_heading(doc,'2.4 网站检索与来源核验门禁',2)
    add_para(doc,'本 Skill 将外部检索从“写作建议”升级为报告生成前的硬门禁：执行 Agent 负责搜索、打开并阅读官方原文；确定性脚本负责校验来源类型、版本绑定、引用解析和冲突状态。')
    ra=data['research_assurance']
    add_table(doc,['门禁项','结果'],[
        ['研究门状态',ra['gate_status']],
        ['Research Plan',ra['research_plan_id']],
        ['Web Research Ledger',ra['research_ledger_id']],
        ['研究时点',ra['research_date']],
        ['检索查询数',ra['queries_executed']],
        ['已打开并阅读的来源',ra['opened_sources']],
        ['已核验一手来源',ra['verified_primary_sources']],
        ['正式政策来源',ra['official_policy_sources']],
        ['正式标准/规范来源',ra['official_standard_sources']],
        ['未解决冲突',ra['unresolved_conflicts']],
        ['门禁回执',ra['gate_report_ref']]
    ],[5.2,11.6],font_size=8.0,first_col_fill=LIGHT_BLUE,cell_margin=42,header_margin=48)
    add_heading(doc,'2.4.1 不可跳过的校验条件',3)
    add_bullet(doc,'必须存在正式政策或主管机关一手来源，并核实文件名称、文号、发布/实施日期、状态和场景边界。')
    add_bullet(doc,'标准或本体的版本、发布日期和用途必须绑定标准组织或官方规范页面；搜索摘要不得作为证据。')
    add_bullet(doc,'报告中的 P/S/R 来源编号必须可解析到 Web Research Ledger；当前事实及关键声明必须由已打开的一手来源支持。')
    add_bullet(doc,'政策日期、场景边界或标准版本存在未解决冲突时，研究门状态必须为 blocked，禁止生成 Word。')
    add_callout(doc,'强制研究门禁',ra.get('notes','生成Word前必须完成网站检索、打开权威原文、核实政策与标准版本，并通过机器门禁。'),fill=LIGHT_ORANGE,edge=ORANGE)

    # 3
    new_page(doc); add_heading(doc,f"3. {data['meta']['domain_name_cn']}领域边界与模块划分",1)
    add_figure(doc,modules,f"图 3-1  {data['meta']['domain_name_cn']}领域模块、共享上层本体与治理依赖")
    add_table(doc,['模块ID','模块','范围','核心对象','依赖'],[[x['id'],x['name'],x['scope'],x['core_objects'],x['dependencies']] for x in data['modules']],[1.5,3.2,4.3,4.5,3.3],font_size=7.9)
    add_heading(doc,'3.1 领域主链',2); add_callout(doc,'主语义链','Situation/Need → Goal/Policy → Service/Product → Process/Case → Task/Decision/Action → Event/Evidence → Outcome/KPI → Episode/Learning',fill=LIGHT_BLUE,edge=NAVY)

    # 4
    new_page(doc); add_heading(doc,'4. 关键参与方与角色',1)
    add_table(doc,['参与方/组织','角色','主要责任','权限或证据'],[[x['party'],x['roles'],x['responsibilities'],x['authority_or_evidence']] for x in data['stakeholders']],[3.0,3.8,6.0,4.0],font_size=8.0)
    add_callout(doc,'角色纪律','Party 与 Role 分离；同一主体可在不同情境扮演不同角色。账号、身份标识和系统记录不是主体本身。',fill=LIGHT_ORANGE,edge=ORANGE)

    # 5
    new_page(doc); add_heading(doc,'5. 主要业务系统与 System of Record',1)
    add_table(doc,['系统','权威数据','典型Source Object','System-of-Record规则'],[[x['system'],x['authoritative_data'],x['source_objects'],x['sor_rule']] for x in data['systems']],[3.5,4.5,4.2,4.6],font_size=7.8)
    add_callout(doc,'SoR原则','分析平台、搜索索引、模型缓存和报表通常是派生视图；正式身份、案件、决定、交易、证据和审计必须回到明确的权威系统。',fill=LIGHT_BLUE,edge=NAVY)

    # 6
    new_page(doc); add_heading(doc,'6. Canonical Concepts 基线',1)
    add_para(doc,f"本轮形成 {len(data['concepts'])} 个候选Canonical Concepts。它们用于暴露共同结构、专业边界和集成点，不代表完整领域本体。")
    groups=defaultdict(list)
    for c in data['concepts']: groups[c.get('group','未分组')].append(c)
    for group,items in groups.items():
        add_heading(doc,f'6.{list(groups.keys()).index(group)+1} {group}',2)
        rows=[[c['id'],c['name_cn'],c['name_en'],c['upper_parent'],c['definition'],joinv(c.get('source_refs',[]),'/'),c['maturity']] for c in items]
        add_table(doc,['ID','中文名','English','Upper Parent','定义','来源','成熟度'],rows,[1.5,2.3,2.8,2.7,5.1,1.4,1.4],font_size=7.3)
    add_heading(doc,'6.x 关键语义混淆',2)
    add_table(doc,['不可混同表达','澄清'],[[x['expression'],x['clarification']] for x in data['semantic_confusions']],[6.4,10.4],font_size=8.5,first_col_fill=LIGHT_ORANGE)

    # 7
    new_page(doc); add_heading(doc,'7. 核心关系三元组',1)
    add_para(doc,'领域本体不能停留在名词清单；每个核心对象必须进入关系、流程、证据、决策或结果链。')
    for idx in range(0,len(data['relations']),25):
        chunk=data['relations'][idx:idx+25]
        add_table(doc,['主语','关系','宾语','业务含义','来源/状态'],[[x['subject'],x['predicate'],x['object'],x['meaning'],x['source_or_status']] for x in chunk],[3.0,2.4,3.0,5.5,2.9],font_size=7.5)

    # 8
    new_page(doc); add_heading(doc,'8. 端到端核心流程',1)
    add_table(doc,['ID','流程','触发','主要阶段'],[[x['id'],x['name'],x['trigger'],x['stages']] for x in data['processes']],[1.6,4.0,4.7,6.5],font_size=8.0)
    add_table(doc,['流程','关键决策','证据','结果'],[[x['name'],x['decisions'],x['evidence'],x['outcome']] for x in data['processes']],[3.7,4.3,4.5,4.3],font_size=7.8)

    # 9
    new_page(doc); add_heading(doc,'9. 事件、决策、规则、风险与 KPI',1)
    add_heading(doc,'9.1 关键事件',2); add_table(doc,['事件','触发','状态变化','证据'],[[x['event'],x['trigger'],x['state_change'],x['evidence']] for x in data['events']],[3.0,4.3,4.6,4.9],font_size=8)
    add_heading(doc,'9.2 关键决策与责任边界',2); add_table(doc,['决策','决定主体','证据','结果','HITL'],[[x['decision'],x['authority'],x['evidence'],x['result'],x['hitl']] for x in data['decisions']],[2.8,3.7,4.3,3.1,2.9],font_size=7.6)
    add_heading(doc,'9.3 规则与约束候选',2); add_table(doc,['规则','REASON','来源/确认状态','优先级'],[[x['rule'],x['reason'],x['source_or_status'],x['priority']] for x in data['rules_constraints']],[4.5,5.4,4.2,2.7],font_size=7.8)
    add_heading(doc,'9.4 风险—控制基线',2); add_table(doc,['风险','后果','控制','责任方','证据'],[[x['risk'],x['consequence'],x['control'],x['owner'],x['evidence']] for x in data['risk_controls']],[2.5,3.8,5.0,2.6,2.9],font_size=7.5)
    add_heading(doc,'9.5 KPI框架',2); add_table(doc,['KPI','定义','度量','来源','注意事项'],[[x['kpi'],x['definition'],x['measure'],x['source'],x['caveat']] for x in data['kpis']],[2.7,4.0,3.3,3.1,3.7],font_size=7.7)

    # 10
    new_page(doc); add_heading(doc,'10. 外部标准、政策语义资产与本体对标',1)
    add_callout(doc,'对标方法','先判定资产角色，再判断复用概念、映射关系、主干适用性、优势、缺口和本地扩展；API/交换对象不能直接代替现实业务实体。',fill=LIGHT_BLUE,edge=NAVY)
    add_heading(doc,'10.1 资产角色与语义主干定位',2)
    add_table(doc,['标准/资产','组织','版本/日期','角色','覆盖','主干?'],[[x['name'],x['organization'],x['version_or_date'],x['role'],x['coverage'],x['backbone_recommendation']] for x in data['standards']],[3.2,2.8,2.6,3.2,3.5,1.5],font_size=7.3)
    new_page(doc); add_heading(doc,'10.2 逐项映射、优势与缺口',2)
    add_table(doc,['标准/资产','复用概念','Mapping','优势','缺口/本地扩展'],[[x['name'],x['reusable_concepts'],x['mapping_relation'],x['strengths'],x['gaps']+'；'+x['local_extension']] for x in data['standards']],[3.0,3.7,2.0,3.6,4.5],font_size=7.2)

    # 11
    new_page(doc); add_heading(doc,'11. Shared Upper Ontology 映射',1)
    add_table(doc,['上层构造','领域概念','使用说明'],[[x['upper_construct'],joinv(x['domain_concepts'],'、'),x['usage']] for x in data['upper_mappings']],[4.0,6.3,6.5],font_size=8.0,first_col_fill=LIGHT_BLUE)

    # 12
    new_page(doc); add_heading(doc,f"12. {data['meta']['domain_name_cn']}智能体扩展、授权、HITL与审计",1)
    ag=data['agent_governance']
    add_heading(doc,'12.1 智能体执行对象',2); add_table(doc,['对象','含义','主要关系'],[[x['name'],x['meaning'],x['relations']] for x in ag['objects']],[4.0,6.2,6.6],font_size=8.2)
    add_heading(doc,'12.2 A0–A4自主性分级',2); add_table(doc,['等级','典型任务','人工模式','约束'],[[x['level'],x['typical_tasks'],x['human_mode'],x['constraints']] for x in ag['autonomy_levels']],[1.5,5.0,4.0,6.3],font_size=8.0,first_col_fill=LIGHT_GREEN)
    add_heading(doc,'12.3 必须由人类最终决定的事项',2)
    for x in ag['human_final_decisions']: add_bullet(doc,x)
    new_page(doc); add_heading(doc,'12.4 Agent Manifest建议字段',2)
    add_para(doc,'Agent Manifest 用于把身份、能力、授权、工具、证据与人工接管边界固化为可审计契约；字段不得只停留在自由文本说明。')
    add_table(doc,['字段'],[[x] for x in ag['manifest_fields']],[16.8],font_size=8.8,cell_margin=50,header_margin=55)
    add_bullet(doc,'身份与版本：agent_id、role、owner/version 用于唯一定位执行主体及其配置版本。')
    add_bullet(doc,'能力与边界：capabilities、allowed_tasks、prohibited_actions、authorization_scope 明确可做、不得做及授权范围。')
    add_bullet(doc,'治理与追溯：required_evidence、hitl_policy、audit_policy、escalation_path 约束证据、人工接管、审计和升级。')
    add_callout(doc,'强治理主链','DomainAgent → Capability → Skill/Tool → Authorization → Candidate Output → Safety/Quality Gate → Human Approval → Action → Audit → Review/Appeal → Episode/Learning',fill=LIGHT_ORANGE,edge=ORANGE)

    # 13
    new_page(doc); add_heading(doc,'13. 闭环示例与 KSTAR/ECS 映射',1)
    add_figure(doc,loop,'图 13-1  可解释业务闭环与KSTAR学习边界')
    for i,x in enumerate(data['closed_loops'],1):
        if i > 2 and i % 2 == 1:
            new_page(doc)
        add_heading(doc,f'13.{i} {x["name"]}',2)
        rows=[('Situation',x['situation']),('Goal',x['goal']),('Process',x['process']),('Task / Decision',x['task_decision']),('Evidence / Event',x['evidence_event']),('Outcome / KPI',x['outcome_kpi']),('Episode / Learning',x['episode_learning'])]
        add_table(doc,['闭环对象','内容'],rows,[4.0,12.8],font_size=8.7,first_col_fill=LIGHT_GREEN)
    add_heading(doc,'13.x KSTAR/ECS六阶段',2); add_table(doc,['阶段','输入','输出','控制'],[[x['stage'],x['inputs'],x['outputs'],x['controls']] for x in data['kstar_mapping']],[1.6,5.1,5.1,5.0],font_size=8.0,first_col_fill=LIGHT_BLUE)

    # 14
    new_page(doc); add_heading(doc,'14. 架构边界、实施建议与待确认问题',1)
    add_heading(doc,'14.1 模块边界与责任',2); add_table(doc,['层/模块','资产','建议责任'],[[x['layer'],x['assets'],x['ownership']] for x in data['architecture_boundaries']],[4.0,7.6,5.2],font_size=8.2)
    add_heading(doc,'14.2 实施优先级',2); add_table(doc,['优先级','工作包','预期结果'],[[x['priority'],x['work_package'],x['result']] for x in data['implementation_priorities']],[1.8,8.0,7.0],font_size=8.3,first_col_fill=LIGHT_ORANGE)
    add_heading(doc,'14.3 待评审/待决策问题',2)
    for q in data['review_questions']: add_number(doc,q)
    add_callout(doc,'汇总建议','共享上层本体、身份/组织/角色、证据/来源、Agent授权审计和KSTAR资产由总体架构统一；专业概念、规则、标准映射和SoR由领域模块负责。',fill=LIGHT_BLUE,edge=NAVY)

    # Appendix A
    new_page(doc); add_heading(doc,'附录 A. Master Mapping Workbook 字段与示例',1)
    add_table(doc,['字段','用途'],[[x['field'],x['purpose']] for x in data['workbook_fields']],[6.0,10.8],font_size=8.5,first_col_fill=LIGHT_BLUE)
    # examples dynamic columns
    keys=[]
    for row in data['workbook_examples']:
        for k in row:
            if k not in keys: keys.append(k)
    if keys:
        width=16.8/len(keys); add_table(doc,keys,[[r.get(k,'') for k in keys] for r in data['workbook_examples']],[width]*len(keys),font_size=max(6.5,8.7-len(keys)*.25))

    # Appendix B
    new_page(doc); add_heading(doc,'附录 B. 来源、版本与证据清单',1)
    add_heading(doc,'B.1 网站检索与来源核验摘要',2)
    ra=data['research_assurance']
    add_para(doc,f"研究门禁={ra['gate_status']}；查询={ra['queries_executed']}；打开来源={ra['opened_sources']}；一手来源={ra['verified_primary_sources']}；正式政策={ra['official_policy_sources']}；正式标准/规范={ra['official_standard_sources']}；未解决冲突={ra['unresolved_conflicts']}。")
    add_para(doc,'完整检索计划、查询记录、页面打开/阅读记录、来源哈希、声明绑定、标准版本核验及冲突处理，见随报告交付的 research-plan.json、web-research-ledger.json 和 research-gate.json。')
    add_heading(doc,'B.2 报告来源台账',2)
    rows=[]
    for s in data['sources']:
        rows.append([s['id'],s.get('evidence_type',''),s['title'],s.get('organization',''),s.get('version_or_date',''),s.get('verification_status',''),s.get('usage',''),s.get('url_or_file','')])
    add_table(doc,['ID','类型','标题','组织','版本/日期','状态','用途','URL/文件'],rows,[1.0,1.9,3.5,2.4,2.1,1.7,2.3,2.0],font_size=6.8)
    add_callout(doc,'来源解释','P/S/I/E/A/R分别代表政策、标准、内部材料、专家决定、分析建议和二手线索。正文来源标记必须可在本附录解析。',fill=LIGHT_GREEN,edge=GREEN)

    # Appendix C
    new_page(doc); add_heading(doc,'附录 C. 来源材料与分析边界说明',1)
    add_table(doc,['材料','处理方式','理由','来源'],[[x['material'],x['treatment'],x['reason'],joinv(x.get('source_refs',[]),'/')] for x in data['material_treatment']],[4.1,5.3,5.5,1.9],font_size=8.0)
    add_heading(doc,'C.1 研究与交付边界',2)
    add_bullet(doc,'本报告为领域本体景观、标准对标和上层映射候选成果，不是完整生产本体或正式行业标准。')
    add_bullet(doc,'外部标准和政策按研究时点记录；版本变化后必须复核。')
    add_bullet(doc,'未经领域专家确认的规则理由、阈值、责任边界和高影响动作保持candidate或pending。')
    add_bullet(doc,'系统表、API对象、JSON字段和文档记录属于系统表示层，不等同于现实业务实体。')
    add_bullet(doc,'自动流程最高到staged，最终接受、发布和生产使用由独立人类流程决定。')

    doc.save(out)
    return {'docx':str(out),'asset_dir':str(asset),'tables':len(doc.tables),'paragraphs':len(doc.paragraphs),'images':3}

def main():
    ap=argparse.ArgumentParser(); ap.add_argument('data'); ap.add_argument('output'); args=ap.parse_args()
    data=load(args.data); result=build(data,args.output); print(json.dumps(result,ensure_ascii=False,indent=2))

if __name__=='__main__': main()
